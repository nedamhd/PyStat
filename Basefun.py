# create table 4 in docx

def py2doc(df, doc=None ,title="", caption="", save=True, filename="Result", path=None):
    import docx
    from os import getcwd, listdir
    #from docx.enum.table import WD_ROW_HEIGHT_RULE
    if (doc==None):
       doc = docx.Document()
    # Initialise the table  
    #doc.add_paragraph(title)
    doc.add_heading(title, 2)

    t = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
    t.style = 'TableGrid'

    # Add the body of the data frame to the table
    for k in range(df.shape[1]):
       cell =  df.columns[k]
       t.cell(0, k).text = str(cell)
    for i in range(1,df.shape[0]+1):
        for j in range(df.shape[1]):
            cell = df.iat[i-1, j]
            t.cell(i, j).text = str(cell)

     # for row in t.rows:  
     #     row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    if caption!="":
        row= t.add_row()
        cell0=row.cells[0]
        for k in range(1,df.shape[1]):
            cell=row.cells[k]
            cell0=cell0.merge(cell)
        row.cells[0].text = str(caption)
    doc.add_paragraph("\n\n")
    if save==True:
       loc=getcwd()
       if path!= None:
          loc=path
       files= listdir(loc)
       i=0
       filenametemp=filename
       while (filenametemp+".docx" in files):
         i=i+1
         filenametemp=filename+str(i)
       doc.save(loc+"\\"+filenametemp+".docx")
       print("The file is saved at " + loc + " as " + filenametemp + ".docx")
    return doc


#Example
from sklearn import datasets
import pandas as pd
from Basefun import py2doc
iris= datasets.load_iris()
df=pd.DataFrame(iris.data).head()
mydoc= py2doc(df=df ,title="International System of Units",save=False)
mydoc= py2doc(df=df, doc=mydoc ,title="International System of Units",save=False)
mydoc= py2doc(df=df, doc=mydoc ,title="International System of Units",save=False)
py2doc(df=df, doc=mydoc ,title="International Sys", caption="Table One")
